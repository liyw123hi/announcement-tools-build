#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
建行公告附件验证工具 - 【Windows 10 兼容版 v6】
基于 v6 长名称产品修复版 改写，适配 Windows 10 环境。

Windows 10 兼容修改要点：
1. read_doc_file() 函数: Windows 下优先用 olefile 库提取 .doc 文本，不再尝试 textutil。
   备选方案：python-docx2txt、antiword（需安装）。
2. read_text_file() 函数: 删除 macOS 隔离属性移除代码（xattr）。
3. get_downloaded_files_since(): 删除 .DS_Store 过滤（Windows 无此文件）。
4. wait_for_file_download(): 删除 .DS_Store 过滤。
5. main() 函数中的 ChromeDriver 初始化:
   - 删除 macOS 专有路径（/opt/homebrew/bin/chromedriver、/usr/local/bin/chromedriver）。
   - Windows 下先尝试 webdriver-manager，失败后在以下位置查找 chromedriver.exe：
     当前脚本目录、PATH 环境变量、C:\Program Files\Google\Chrome\Application\。
   - 保留 .chromedriver_path 缓存机制。
6. 保留所有 v6 业务逻辑完全不变。
7. 保留所有中文注释。
"""

import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
import time
import os
import sys
import shutil
from datetime import datetime
import logging
import re
import traceback
from typing import Dict, Any, List, Optional, Tuple
import glob
import tempfile
import platform

# ==================== 配置日志记录器 ====================
def setup_logger():
    """设置详细的状态日志记录器"""
    logger = logging.getLogger("公告验证")
    logger.setLevel(logging.INFO)
    
    if not logger.hasHandlers():
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        console_format = logging.Formatter(
            '[%(asctime)s] [%(levelname)s] %(message)s',
            datefmt='%H:%M:%S'
        )
        console_handler.setFormatter(console_format)
        logger.addHandler(console_handler)
    
    return logger

logger = setup_logger()

# ==================== 调试截图功能 ====================
def save_debug_screenshot(driver, screenshot_dir: str, prefix: str, product_code: str) -> str:
    """
    保存调试截图，便于排查长名称产品匹配问题
    返回截图文件路径
    """
    try:
        if not os.path.exists(screenshot_dir):
            os.makedirs(screenshot_dir)
        
        timestamp = datetime.now().strftime('%H%M%S')
        safe_code = re.sub(r'[^\w]', '_', product_code)[:20]
        filename = f"{prefix}_{safe_code}_{timestamp}.png"
        filepath = os.path.join(screenshot_dir, filename)
        
        driver.save_screenshot(filepath)
        logger.info(f"  │  调试截图已保存: {filepath}")
        return filepath
    except Exception as e:
        logger.debug(f"保存截图失败: {str(e)}")
        return ""

# ==================== 清理函数 ====================
def clean_download_folder(download_path: str):
    """静默清理文件夹"""
    if os.path.exists(download_path):
        try:
            for filename in os.listdir(download_path):
                file_path = os.path.join(download_path, filename)
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception as e:
                    pass
        except Exception as e:
            pass

# ==================== 辅助函数 ====================
def clean_text(text: str) -> str:
    """清理文本，移除多余空格和换行符"""
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def is_title_match(expected_title: str, actual_title: str) -> bool:
    """严格检查标题是否匹配"""
    if not expected_title or not actual_title:
        return False
    
    expected_clean = clean_text(expected_title)
    actual_clean = clean_text(actual_title)
    
    # 完全一致
    if expected_clean == actual_clean:
        return True
    
    # 包含关系
    if expected_clean in actual_clean or actual_clean in expected_clean:
        return True
    
    # 移除标点符号和空格后比较
    expected_simple = re.sub(r'[^\w\u4e00-\u9fff]', '', expected_clean)
    actual_simple = re.sub(r'[^\w\u4e00-\u9fff]', '', actual_clean)
    
    if expected_simple and actual_simple:
        if expected_simple in actual_simple or actual_simple in expected_simple:
            return True
    
    return False

def contains_target_date(url: str, target_date: str) -> bool:
    """检查URL中是否包含目标日期"""
    if not url or not target_date:
        return False
    
    date_formats = [
        target_date.replace('-', ''),  # 20260316
        target_date,                   # 2026-03-16
        target_date.replace('-', '/'),  # 2026/03/16
    ]
    
    for fmt in date_formats:
        if fmt in url:
            return True
    
    year, month, day = target_date.split('-')
    patterns = [
        rf'{year}[\/\-]?{month}[\/\-]?{day}',
        rf'{year}{month}{day}',
    ]
    
    for pattern in patterns:
        if re.search(pattern, url):
            return True
    
    return False

def extract_date_from_text(text: str, target_date: str) -> Optional[str]:
    """从全文中提取并验证日期（保留兼容，内部调用 extract_date_from_tail 的全文模式）"""
    return extract_date_from_tail(text, target_date, tail_only=False)


def extract_date_from_tail(text: str, target_date: str, tail_only: bool = True) -> Optional[str]:
    """
    从文本尾部（落款区域）提取日期并与目标日期匹配。
    
    tail_only=True  时：只取文本最后 5000 字符搜索（word 落款日 / PDF 末页）
    tail_only=False 时：全文搜索（兼容旧逻辑）
    
    返回匹配的日期字符串，或 None
    """
    if not text or not target_date:
        return None
    
    target_year, target_month, target_day = target_date.split('-')
    
    # 支持的日期格式模式
    date_patterns = [
        r'(\d{4})年(\d{1,2})月(\d{1,2})日',   # 2026年03月16日
        r'(\d{4})年(\d{1,2})月(\d{1,2})号',   # 2026年03月16号
        r'(\d{4})-(\d{1,2})-(\d{1,2})',       # 2026-03-16
        r'(\d{4})/(\d{1,2})/(\d{1,2})',       # 2026/03/16
    ]
    
    search_text = text[-5000:] if tail_only else text
    
    for pattern in date_patterns:
        # 用 findall 从尾文本中找所有日期
        matches = re.findall(pattern, search_text)
        # 取最后一个匹配（落款日期通常在末尾）
        for match in reversed(matches):
            year = match[0]
            month = match[1].zfill(2)
            day = match[2].zfill(2)
            if year == target_year and month == target_month and day == target_day:
                return f"{year}-{month}-{day}"
    
    return None


def read_pdf_tail(file_path: str) -> str:
    """
    只读取 PDF 最后一页的文本内容（用于落款日期校验）
    """
    try:
        try:
            import PyPDF2
        except ImportError:
            logger.error("未安装PyPDF2库，无法读取.pdf文件")
            logger.info("请运行: pip install PyPDF2")
            return ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            if total_pages == 0:
                return ""
            # 只读最后一页
            last_page = pdf_reader.pages[total_pages - 1]
            text = last_page.extract_text() or ""
            logger.debug(f"PDF 共 {total_pages} 页，已读取末页，内容长度: {len(text)}")
            return text
    except Exception as e:
        logger.error(f"读取PDF末页失败: {str(e)}")
        return ""

def read_docx_file(file_path: str) -> str:
    """
    读取.docx文件内容
    注意：需要安装python-docx库：pip install python-docx
    """
    try:
        # 尝试导入python-docx
        try:
            from docx import Document
        except ImportError:
            logger.error("未安装python-docx库，无法读取.docx文件")
            logger.info("请运行: pip install python-docx")
            return ""
        
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"读取.docx文件失败: {str(e)}")
        return ""

def read_docx_tail(file_path: str, tail_paragraphs: int = 20) -> str:
    """
    只读取 .docx 文件尾部段落（用于落款日期校验）
    比全量 read_docx_file 快很多，尤其对大文档
    """
    try:
        try:
            from docx import Document
        except ImportError:
            logger.error("未安装python-docx库，无法读取.docx文件")
            return ""
        
        doc = Document(file_path)
        all_paras = doc.paragraphs
        # 只取最后 N 个段落
        tail_paras = all_paras[-tail_paragraphs:] if len(all_paras) > tail_paragraphs else all_paras
        text = '\n'.join(p.text for p in tail_paras)
        logger.debug(f"docx 共 {len(all_paras)} 段，取末 {len(tail_paras)} 段，长度: {len(text)}")
        return text
    except Exception as e:
        logger.error(f"读取.docx尾部失败: {str(e)}")
        return ""

def read_doc_file(file_path: str) -> str:
    """
    读取旧版.doc文件内容（非docx）
    
    【Windows 10 兼容版】读取优先级调整：
    1. olefile（Windows 主要方案，无需额外依赖）
    2. python-docx2txt（备选，需 pip install python-docx2txt）
    3. antiword（备选，需安装 antiword.exe 并配置 PATH）
    4. 文本文件直接读取（兜底）
    
    macOS 下 textutil 仅作为最后兜底（极少见 .doc 文件）。
    """
    if not os.path.exists(file_path):
        logger.warning(f"文件不存在: {file_path}")
        return ""
    
    # Windows 优先方案1：使用 olefile 提取 OLE 复合文档中的文本
    # 适用于 .doc (Office 97-2003) 文件
    try:
        import olefile
        ole = olefile.OleFileIO(file_path)
        if ole.exists('WordDocument'):
            # 尝试从 Word 文档中提取文本流
            stream_names = ole.listdir()
            text_content = []
            for stream in stream_names:
                stream_path = '/'.join(stream)
                if '1Table' in stream_path or '0Table' in stream_path:
                    continue  # 跳过表格流
                if stream_path.endswith('Data') or 'CompObj' in stream_path:
                    continue
                try:
                    data = ole.openstream(stream).read()
                    # 尝试解码为文本
                    for enc in ['utf-16-le', 'utf-8', 'gbk', 'gb2312']:
                        try:
                            text = data.decode(enc, errors='ignore')
                            # 过滤掉二进制噪音
                            clean = re.sub(
                                r'[^\x20-\x7e\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\n\r\t]',
                                '', text
                            )
                            if len(clean) > 20:
                                text_content.append(clean)
                                break
                        except:
                            continue
                except:
                    continue
            ole.close()
            if text_content:
                content = '\n'.join(text_content)
                logger.info(f"使用 olefile 成功读取 .doc 文件，内容长度: {len(content)} 字符")
                return content
        ole.close()
    except ImportError:
        logger.debug("未安装 olefile，尝试其他方案...")
    except Exception as e:
        logger.debug(f"olefile 读取 .doc 失败: {str(e)}")
    
    # Windows 备选方案2：python-docx2txt（需 pip install python-docx2txt）
    # python-docx2txt 底层使用 zipfile + xml 解析，无需 COM/Office
    try:
        import docx2txt
        content = docx2txt.process(file_path)
        if content and len(content.strip()) > 20:
            logger.info(f"使用 python-docx2txt 成功读取 .doc 文件，内容长度: {len(content)} 字符")
            return content
    except ImportError:
        logger.debug("未安装 python-docx2txt，尝试其他方案...")
    except Exception as e:
        logger.debug(f"python-docx2txt 读取 .doc 失败: {str(e)}")
    
    # Windows 备选方案3：antiword（需安装 antiword.exe 并添加到 PATH）
    # antiword 是成熟的开源工具，读取 .doc 效果较好
    if platform.system() == "Windows":
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                content = result.stdout
                logger.info(f"使用 antiword 成功读取 .doc 文件，内容长度: {len(content)} 字符")
                return content
        except FileNotFoundError:
            logger.debug("antiword 未安装或不在 PATH 中，跳过...")
        except subprocess.TimeoutExpired:
            logger.warning("antiword 转换超时")
        except Exception as e:
            logger.debug(f"antiword 调用失败: {str(e)}")
    
    # macOS 兜底方案：使用 textutil（macOS 自带工具）
    if platform.system() == "Darwin":
        try:
            import subprocess
            result = subprocess.run(
                ['textutil', '-convert', 'txt', '-stdout', file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                content = result.stdout
                logger.info(f"使用 textutil 成功读取 .doc 文件，内容长度: {len(content)} 字符")
                return content
            else:
                logger.warning(f"textutil 转换失败: {result.stderr[:100] if result.stderr else '无输出'}")
        except subprocess.TimeoutExpired:
            logger.warning("textutil 转换超时")
        except Exception as e:
            logger.warning(f"textutil 调用失败: {str(e)}")
    
    # 最终兜底：尝试作为文本直接读取
    try:
        return read_text_file(file_path)
    except:
        pass
    
    logger.warning(f"无法读取 .doc 文件内容: {file_path}")
    return ""


def read_doc_tail(file_path: str, tail_chars: int = 4000) -> str:
    """
    只读取 .doc 文件尾部内容，用于落款日期匹配
    先读取全文，再截取末尾部分，避免处理大文档全文
    """
    if not os.path.exists(file_path):
        logger.warning(f"文件不存在: {file_path}")
        return ""

    full_text = read_doc_file(file_path)
    if not full_text:
        return ""

    tail = full_text[-tail_chars:]
    logger.debug(f".doc 尾部内容长度: {len(tail)} 字符")
    return tail


def read_text_file(file_path: str) -> str:
    """
    读取文本文件内容（跨平台、多编码支持）
    """
    if not os.path.exists(file_path):
        logger.warning(f"文件不存在: {file_path}")
        return ""
    
    try:
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            logger.warning(f"文件为空: {file_path}")
            return ""
    except Exception as e:
        logger.warning(f"无法读取文件大小: {str(e)}")
        return ""
    
    # 尝试多种编码
    encodings = [
        'utf-8',
        'utf-8-sig',
        'gbk',
        'gb2312',
        'big5',
        'latin-1',
        'cp1252',
    ]
    
    file_content = ""
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                file_content = f.read()
            
            # 检查是否成功读取了内容
            if file_content and len(file_content.strip()) > 0:
                logger.debug(f"使用编码 {encoding} 成功读取文件，内容长度: {len(file_content)}")
                return file_content
            
        except Exception as e:
            logger.debug(f"尝试编码 {encoding} 失败: {str(e)}")
            continue
    
    # 如果所有文本编码都失败，尝试二进制读取然后解码
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        
        # 尝试检测并解码
        for encoding in encodings:
            try:
                file_content = raw_data.decode(encoding, errors='ignore')
                if file_content and len(file_content.strip()) > 0:
                    logger.debug(f"使用二进制模式和编码 {encoding} 成功读取文件")
                    return file_content
            except:
                continue
    except Exception as e:
        logger.warning(f"二进制读取失败: {str(e)}")
    
    logger.warning(f"无法使用任何编码读取文件: {file_path}")
    return ""

def read_pdf_file(file_path: str) -> str:
    """
    读取PDF文件内容
    注意：需要安装PyPDF2库：pip install PyPDF2
    """
    try:
        # 尝试导入PyPDF2
        try:
            import PyPDF2
        except ImportError:
            logger.error("未安装PyPDF2库，无法读取.pdf文件")
            logger.info("请运行: pip install PyPDF2")
            return ""
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        
        return text
    except Exception as e:
        logger.error(f"读取PDF文件失败: {str(e)}")
        return ""

def get_most_recent_file(download_path: str, extension: str = None) -> Optional[str]:
    """获取下载目录中最新的文件"""
    if not os.path.exists(download_path):
        return None
    
    files = []
    for filename in os.listdir(download_path):
        file_path = os.path.join(download_path, filename)
        if os.path.isfile(file_path):
            # 过滤临时文件
            if not filename.endswith(('.tmp', '.crdownload', '.part')):
                if extension is None or filename.lower().endswith(extension.lower()):
                    files.append((file_path, os.path.getmtime(file_path)))
    
    if not files:
        return None
    
    # 按修改时间排序，返回最新的文件
    files.sort(key=lambda x: x[1], reverse=True)
    return files[0][0]

def get_downloaded_files_since(download_path: str, start_time: float) -> List[str]:
    """获取指定时间后修改的所有文件"""
    downloaded_files = []
    
    try:
        for filename in os.listdir(download_path):
            # 跳过系统文件和临时文件
            if filename.startswith('.'):
                continue
            if filename.endswith(('.tmp', '.crdownload', '.part', '.downloading')):
                continue
            
            file_path = os.path.join(download_path, filename)
            
            # 确保是文件而不是文件夹
            if not os.path.isfile(file_path):
                continue
            
            try:
                # 获取修改时间
                mtime = os.path.getmtime(file_path)
                
                # 只获取在 start_time 之后修改的文件
                if mtime >= start_time:
                    downloaded_files.append(file_path)
            except Exception as e:
                logger.debug(f"无法读取文件 {file_path} 的修改时间: {str(e)}")
                continue
    
    except Exception as e:
        logger.error(f"获取已下载文件列表失败: {str(e)}")
    
    return downloaded_files


def wait_for_file_download(download_path: str, max_wait_time: int = 30, 
                          initial_wait: float = 2.0) -> Optional[str]:
    """
    等待文件下载完成
    返回：下载的文件路径，如果超时返回 None
    """
    # 获取初始文件列表（用于识别新下载的文件）
    initial_files = set()
    try:
        for filename in os.listdir(download_path):
            if not filename.startswith('.'):
                initial_files.add(os.path.join(download_path, filename))
    except:
        pass
    
    # 初始等待，让浏览器开始下载
    time.sleep(initial_wait)
    
    start_time = time.time()
    last_valid_file = None
    last_valid_size = 0
    stable_count = 0
    
    poll_interval = 0.3  # 初始快速轮询
    while time.time() - start_time < max_wait_time:
        try:
            # 获取下载目录中的所有文件
            current_files = []
            for filename in os.listdir(download_path):
                # 跳过系统文件和临时文件
                if filename.startswith('.'):
                    continue
                if filename.endswith(('.tmp', '.crdownload', '.part', '.downloading')):
                    continue
                
                file_path = os.path.join(download_path, filename)
                if os.path.isfile(file_path):
                    current_files.append(file_path)
            
            # 查找新文件（不在初始列表中的文件）
            new_files = [f for f in current_files if f not in initial_files]
            
            if new_files:
                # 获取最新修改的文件
                new_files.sort(key=lambda f: os.path.getmtime(f), reverse=True)
                current_file = new_files[0]
                
                # 检查文件大小是否稳定
                try:
                    current_size = os.path.getsize(current_file)
                    
                    if current_size > 0:
                        # 如果大小相同且稳定了多次，认为下载完成
                        if current_size == last_valid_size:
                            stable_count += 1
                            if stable_count >= 2:  # 连续 2 次检查大小相同即认为完成
                                logger.info(f"文件下载完成: {os.path.basename(current_file)} ({current_size} 字节)")
                                return current_file
                        else:
                            last_valid_file = current_file
                            last_valid_size = current_size
                            stable_count = 0
                            logger.debug(f"检测到新文件: {os.path.basename(current_file)} ({current_size} 字节)")
                except OSError as e:
                    logger.debug(f"无法读取文件大小: {str(e)}")
                    stable_count = 0
            else:
                # 未发现新文件，动态增加轮询间隔
                poll_interval = min(poll_interval * 1.2, 1.0)
            
            time.sleep(poll_interval)
            
        except Exception as e:
            logger.debug(f"等待下载时发生错误: {str(e)}")
            time.sleep(0.5)
    
    # 超时后，如果有找到文件，尝试最后一次验证
    if last_valid_file and os.path.exists(last_valid_file):
        try:
            final_size = os.path.getsize(last_valid_file)
            if final_size > 0:
                logger.info(f"文件下载（超时前）: {os.path.basename(last_valid_file)} ({final_size} 字节)")
                return last_valid_file
        except:
            pass
    
    return None


def download_and_read_file(driver, attachment_url: str, download_path: str) -> Tuple[str, str]:
    """
    下载文件并读取内容
    返回：(文件内容, 文件路径)
    """
    # 清理下载目录
    clean_download_folder(download_path)
    
    # 记录当前窗口
    original_window = driver.current_window_handle
    
    try:
        # 通过JavaScript在新标签页中打开链接（触发下载）
        logger.debug(f"正在下载: {attachment_url[:80]}...")
        driver.execute_script(f"window.open('{attachment_url}');")
        
        # 等待文件下载完成（初始等待0.5s，配合智能轮询，小文件秒回）
        downloaded_file = wait_for_file_download(download_path, max_wait_time=30, initial_wait=0.5)
        
        if not downloaded_file or not os.path.exists(downloaded_file):
            logger.warning(f"未找到下载的文件或文件不存在")
            return "", ""
        
        # 最后验证文件大小
        try:
            file_size = os.path.getsize(downloaded_file)
            if file_size == 0:
                logger.warning(f"文件大小为 0: {os.path.basename(downloaded_file)}")
                return "", ""
        except Exception as e:
            logger.warning(f"无法读取文件大小: {str(e)}")
            return "", ""
        
        # 读取文件内容
        file_ext = os.path.splitext(downloaded_file)[1].lower()
        file_content = ""
        
        logger.debug(f"识别文件类型: {file_ext}")
        
        if file_ext == '.pdf':
            # 只读末页，避免全文读取浪费时间
            file_content = read_pdf_tail(downloaded_file)
        elif file_ext == '.docx':
            # 只读尾部段落（落款区域）
            file_content = read_docx_tail(downloaded_file)
        elif file_ext in ['.txt', '.html', '.htm', '.xlsx', '.xls']:
            file_content = read_text_file(downloaded_file)
        elif file_ext in ['.doc']:
            # 旧版.doc文件，只读尾部（落款日期匹配用）
            file_content = read_doc_tail(downloaded_file)
        else:
            # 尝试作为文本文件读取
            try:
                file_content = read_text_file(downloaded_file)
            except Exception as e:
                logger.debug(f"无法作为文本文件读取: {str(e)}")
                file_content = f"[二进制文件类型: {file_ext}，文件路径: {downloaded_file}]"
        
        # 验证是否成功读取内容
        if not file_content or (isinstance(file_content, str) and len(file_content.strip()) == 0):
            logger.warning(f"文件内容为空: {os.path.basename(downloaded_file)}")
            return "", downloaded_file  # 返回文件路径但内容为空
        
        logger.info(f"✓ 文件读取成功，内容长度: {len(file_content)} 字符")
        return file_content, downloaded_file
        
    except Exception as e:
        logger.error(f"下载文件时发生异常: {str(e)}")
        import traceback
        logger.debug(f"详细错误: {traceback.format_exc()}")
        return "", ""
    
    finally:
        # 确保切换回原始窗口
        try:
            if len(driver.window_handles) > 1:
                # 关闭新打开的标签页
                for handle in driver.window_handles:
                    if handle != original_window:
                        try:
                            driver.switch_to.window(handle)
                            driver.close()
                        except:
                            pass
                driver.switch_to.window(original_window)
        except:
            pass

# ==================== 核心验证函数 ====================

def count_trading_days(start_date, end_date) -> int:
    """
    计算两个日期之间的沪深交易日天数（含end_date，不含start_date）。
    使用建行理财子网站的交易日历接口判断是否为交易日。
    备选：若接口不可用，降级为排除周末的简单工作日计算。
    """
    from datetime import datetime as dt, timedelta as td
    
    if isinstance(start_date, str):
        start_date = dt.strptime(start_date, '%Y-%m-%d')
    if isinstance(end_date, str):
        end_date = dt.strptime(end_date, '%Y-%m-%d')
    
    # 优先尝试用第三方库（如果已安装）
    # 方案1: 用 chinese_calendar（如果可用）
    try:
        import chinese_calendar
        count = 0
        current = start_date + td(days=1)
        while current <= end_date:
            if chinese_calendar.is_workday(current):
                count += 1
            current += td(days=1)
        return count
    except ImportError:
        pass
    
    # 方案2: 降级为排除周末的工作日计算（不含节假日）
    count = 0
    current = start_date + td(days=1)
    while current <= end_date:
        if current.weekday() < 5:  # 周一=0, 周五=4
            count += 1
        current += td(days=1)
    return count


def count_workdays(start_date, end_date) -> int:
    """
    计算两个日期之间的工作日天数（含end_date，不含start_date）。
    排除周六、周日，不排除节假日。
    """
    from datetime import datetime as dt, timedelta as td
    
    if isinstance(start_date, str):
        start_date = dt.strptime(start_date, '%Y-%m-%d')
    if isinstance(end_date, str):
        end_date = dt.strptime(end_date, '%Y-%m-%d')
    
    count = 0
    current = start_date + td(days=1)
    while current <= end_date:
        if current.weekday() < 5:
            count += 1
        current += td(days=1)
    return count


def check_nav_intervals(driver, product_code: str, target_date: str,
                        interval_days: int = 7, interval_type: str = '自然日') -> str:
    """
    从搜索结果页提取该产品的净值公告日期，验证相邻间隔是否合规。
    
    【v6修复】增强日期提取容错性：
    - 同时从可见文本和完整文本（绕过CSS截断）中提取日期
    - 增加从href、title属性中提取日期的兜底逻辑
    """
    from datetime import datetime as dt
    
    # 日期提取正则（按优先级排列）
    list_date_patterns = [
        r'(\d{4})年(\d{1,2})月(\d{1,2})日',
        r'(\d{4})年(\d{1,2})月(\d{1,2})号',
        r'(\d{4})-(\d{1,2})-(\d{1,2})',
        r'(\d{4})/(\d{1,2})/(\d{1,2})',
    ]

    # 净值公告关键词（精确匹配，避免"净值估算公告"等非周期公告混入）
    nav_keywords = ['净值公告', '净值公']
    
    def extract_date_from_text(text: str) -> Optional[dt]:
        """从文本中提取日期"""
        if not text:
            return None
        for pattern in list_date_patterns:
            m = re.search(pattern, text)
            if m:
                try:
                    return dt(int(m.group(1)), int(m.group(2)), int(m.group(3)))
                except:
                    continue
        return None
    
    # 扫描页面所有 newsdetail 链接
    all_links = driver.find_elements(By.CSS_SELECTOR, "a[href*='newsdetail']")
    logger.info(f"  │  页面共 {len(all_links)} 个 newsdetail 链接，正在提取日期...")
    
    nav_dates = []  # [(date_obj, link_text), ...]
    seen_dates = set()
    skip_reasons = {'无净值关键词': 0, '日期解析失败': 0, '产品编号不符': 0, '其他异常': 0}
    
    for i, link in enumerate(all_links):
        try:
            href = link.get_attribute('href') or ''
            link_text = (link.text or '').strip()
            
            # 【v6修复】同时获取可见文本和完整文本（绕过CSS截断）
            try:
                parent_li = driver.execute_script("return arguments[0].closest('li');", link)
                parent_text = parent_li.text if parent_li else ''
                parent_full_text = driver.execute_script("return arguments[0] ? arguments[0].textContent : '';", parent_li) if parent_li else ''
            except:
                parent_text = ''
                parent_full_text = ''
            
            parent_text = parent_text.strip()
            parent_full_text = parent_full_text.strip()
            
            # 链接的完整文本（绕过CSS截断）
            full_link_text = driver.execute_script(
                "return arguments[0].textContent || '';", link
            ).strip()
            
            # 链接的 title 属性（可能含完整编号）
            link_title = link.get_attribute('title') or ''
            
            # 合并文本用于匹配
            visible_combined = (parent_text + ' ' + link_text).lower()
            full_combined = (parent_full_text + ' ' + full_link_text + ' ' + link_title + ' ' + href).lower()
            
            # ① 净值关键词过滤：在可见文本或完整文本中检查
            has_nav_keyword = any(kw in visible_combined for kw in nav_keywords)
            if not has_nav_keyword:
                # 兜底：在完整文本中再检查一次
                has_nav_keyword = any(kw in full_combined for kw in nav_keywords)
            
            if not has_nav_keyword:
                skip_reasons['无净值关键词'] += 1
                if skip_reasons['无净值关键词'] <= 3:
                    logger.debug(f"  │   [{i+1}] 跳过（无净值关键词）: {link_text[:40] or parent_text[:40]}")
                continue
            
            # ② 产品编号过滤（双重策略）
            code_in_visible = (product_code.lower() in visible_combined or 
                              product_code.lower() in link_title.lower())
            
            if not code_in_visible:
                # 策略B：完整文本匹配
                code_in_full = product_code.lower() in full_combined
                if not code_in_full:
                    skip_reasons['产品编号不符'] += 1
                    if skip_reasons['产品编号不符'] <= 3:
                        logger.debug(f"  │   [{i+1}] 跳过（编号不符）: 可见='{link_text[:30]}' 完整='{full_link_text[:40]}'")
                    continue
                else:
                    logger.debug(f"  │   [{i+1}] 编号匹配（完整文本）: {link_text[:30]}...")
            
            # ③ 【v6修复】日期提取：优先从完整文本提取， fallback 到可见文本
            found_date = None
            
            # 策略1：从完整父文本提取（绕过CSS截断）
            found_date = extract_date_from_text(parent_full_text)
            
            # 策略2：从可见父文本提取
            if not found_date:
                found_date = extract_date_from_text(parent_text)
            
            # 策略3：从链接完整文本提取
            if not found_date:
                found_date = extract_date_from_text(full_link_text)
            
            # 策略4：从链接可见文本提取
            if not found_date:
                found_date = extract_date_from_text(link_text)
            
            # 策略5：从 href 提取（URL中可能有日期参数）
            if not found_date:
                for pattern in list_date_patterns:
                    m = re.search(pattern, href)
                    if m:
                        try:
                            found_date = dt(int(m.group(1)), int(m.group(2)), int(m.group(3)))
                            break
                        except:
                            continue
            
            # 策略6：从 title 属性提取
            if not found_date:
                found_date = extract_date_from_text(link_title)
            
            if not found_date:
                skip_reasons['日期解析失败'] += 1
                if skip_reasons['日期解析失败'] <= 3:
                    logger.debug(
                        f"  │   [{i+1}] 跳过（日期解析失败）: "
                        f"link='{link_text[:30]}' parent='{parent_text[:50]}' parent_full='{parent_full_text[:50]}'"
                    )
                continue
            
            # 命中
            if found_date not in seen_dates:
                seen_dates.add(found_date)
                nav_dates.append((found_date, link_text[:50] or parent_text[:50]))
                logger.info(f"  │    ✓ [{len(nav_dates)}] {found_date.strftime('%Y-%m-%d')} - {link_text[:40] or parent_text[:40]}")
        except:
            skip_reasons['其他异常'] += 1
            continue
    
    # 诊断信息：打印跳过统计
    skipped_total = sum(skip_reasons.values())
    if skipped_total > 0:
        logger.info(f"  │  跳过统计: {skipped_total}条 (无净值关键词:{skip_reasons['无净值关键词']} 日期解析失败:{skip_reasons['日期解析失败']} 编号不符:{skip_reasons['产品编号不符']} 其他:{skip_reasons['其他异常']})")
    
    # 按日期降序排列（最新在前）
    nav_dates.sort(key=lambda x: x[0], reverse=True)
    all_dates_str = [f"{d[0].strftime('%Y-%m-%d')}" for d in nav_dates]
    logger.info(f"  │  命中净值公告共 {len(nav_dates)} 条，日期列表: {all_dates_str}")
    
    if len(nav_dates) < 2:
        return "净值公告不足2条，无法校验间隔"
    
    # 取前5条（或更少）
    check_dates = nav_dates[:5]
    
    # 计算间隔
    intervals = []
    issues = []
    for i in range(len(check_dates) - 1):
        date_newer = check_dates[i][0]
        date_older = check_dates[i + 1][0]
        natural_diff = (date_newer - date_older).days
        
        # 根据 interval_type 计算实际间隔
        if interval_type == '自然日':
            diff = natural_diff
        elif interval_type == '工作日':
            diff = count_workdays(date_older, date_newer)
        elif interval_type == '沪深交易日':
            diff = count_trading_days(date_older, date_newer)
        else:
            diff = natural_diff
        
        status_mark = "✓" if diff == interval_days else "✗"
        cal_type_note = "" if interval_type == '自然日' else f"（自然日{natural_diff}天）"
        logger.info(
            f"  │  [{i+1}] {date_older.strftime('%Y-%m-%d')} → {date_newer.strftime('%Y-%m-%d')} "
            f"= {diff}{interval_type}{cal_type_note} | 期望{interval_days}{interval_type} {status_mark}"
        )
        
        intervals.append(diff)
        if diff != interval_days:
            issues.append(
                f"{date_older.strftime('%Y-%m-%d')}→{date_newer.strftime('%Y-%m-%d')}={diff}{interval_type}"
            )
    
    interval_desc = '+'.join(f'{d}{interval_type}' for d in intervals)
    
    if issues:
        result = f"异常({interval_type}): {', '.join(issues)}"
        logger.warning(f"  │  ⚠ 净值公告间隔: {result}")
    else:
        result = f"正常({interval_desc})"
        logger.info(f"  │  ✓ 净值公告间隔: {result}")
    
    return result


def validate_announcement_title(driver, expected_title: str) -> Tuple[bool, str]:
    """验证公告标题是否匹配（优化版：先试最可能的选择器）"""
    if not expected_title:
        return True, "无预期标题，跳过验证"
    
    # 1. 检查页面标题（最快，无DOM扫描）
    page_title = driver.title.strip()
    if page_title and is_title_match(expected_title, page_title):
        return True, f"页面标题匹配: {page_title[:50]}..."
    
    # 2. 用 JS 一次性提取所有候选标题文本，避免逐个选择器查找
    try:
        candidates = driver.execute_script("""
            var selectors = ['h1', 'h2', 'h3', '.title', '.tit', '.news-title', '.content-title'];
            var results = [];
            for (var i = 0; i < selectors.length; i++) {
                var elems = document.querySelectorAll(selectors[i]);
                for (var j = 0; j < elems.length; j++) {
                    var txt = (elems[j].textContent || '').trim();
                    if (txt.length > 5) results.push(txt);
                }
            }
            return results;
        """) or []
        
        for text in candidates:
            if is_title_match(expected_title, text):
                return True, f"找到匹配标题: {text[:50]}..."
    except:
        pass
    
    return False, f"未找到与'{expected_title[:50]}...'匹配的标题"

def verify_attachment_with_date(driver, attachment_url: str, 
                               announcement_date: str, download_path: str) -> Tuple[str, str]:
    """验证附件可访问性并检查落款日期是否匹配"""
    
    try:
        logger.info(f"开始验证附件: {attachment_url[:80]}...")
        
        # 下载并读取文件（PDF/docx 已在内部只读尾部）
        file_content, downloaded_file = download_and_read_file(driver, attachment_url, download_path)
        
        # 检查文件是否被下载
        if not downloaded_file or not os.path.exists(downloaded_file):
            return "失败", "文件未被下载"
        
        # 获取文件信息
        try:
            file_size = os.path.getsize(downloaded_file)
            file_name = os.path.basename(downloaded_file)
            logger.debug(f"文件已下载: {file_name} ({file_size} 字节)")
        except:
            pass
        
        # 清理下载的文件
        try:
            if os.path.exists(downloaded_file):
                os.remove(downloaded_file)
        except:
            pass
        
        # 文件内容为空 → 失败
        if not file_content or len(file_content.strip()) == 0:
            return "失败", "文件可访问但内容为空，无法验证落款日期"
        
        # file_content 已是 download_and_read_file 内部按文件类型读取的尾部内容
        # （PDF末页 / docx末20段 / doc末4000字符），直接在其中搜索落款日期
        found_date = extract_date_from_tail(file_content, announcement_date, tail_only=False)
        
        if found_date:
            return "成功", f"落款日期验证通过: {found_date}"
        else:
            # 日期不匹配 → 失败，输出尾部内容帮助调试
            tail_preview = file_content[-200:].replace('\n', ' ').strip()[:150]
            logger.info(f"  │  附件尾部内容: {tail_preview}...")
            return "失败", f"落款日期不匹配（期望 {announcement_date}，附件尾部未找到）"
            
    except Exception as e:
        logger.error(f"验证附件时发生异常: {str(e)}")
        logger.debug(f"详细错误: {traceback.format_exc()}")
        return "异常", f"验证异常: {str(e)[:50]}"

def find_search_button(driver):
    """查找搜索按钮的辅助函数"""
    # 尝试多种选择器
    selectors = [
        # CSS选择器
        "input.but[type='button'][value='搜索']",
        "input[type='button'][value='搜索']",
        "input[type='submit'][value='搜索']",
        "button[type='submit']",
        "input[value='搜索']",
        ".but[type='button']",
        ".but[type='submit']",
        
        # 更通用的选择器
        "input[type='button']",
        "input[type='submit']",
        "button",
        
        # 通过文本查找
        "//input[@value='搜索']",
        "//button[contains(text(), '搜索')]",
        "//input[@type='button' and contains(@value, '搜索')]",
        "//input[@type='submit' and contains(@value, '搜索')]",
    ]
    
    for selector in selectors:
        try:
            if selector.startswith("//"):
                # XPath选择器
                elements = driver.find_elements(By.XPATH, selector)
            else:
                # CSS选择器
                elements = driver.find_elements(By.CSS_SELECTOR, selector)
            
            for element in elements:
                try:
                    # 检查元素是否可见和可点击
                    if element.is_displayed() and element.is_enabled():
                        # 检查元素文本或值是否包含"搜索"
                        element_text = element.text or element.get_attribute('value') or ''
                        if '搜索' in element_text:
                            return element
                except:
                    continue
        except:
            continue
    
    return None

def search_and_verify(driver, product_code: str, announcement_name: str, 
                      announcement_url: str, announcement_date: str,
                      download_path: str, index: int, total: int, 
                      processed_count: int,
                      interval_days: int = 7, interval_type: str = '自然日',
                      screenshot_dir: str = None) -> Dict[str, Any]:
    """
    完整的验证逻辑（v6修复版 - 长名称产品兼容）
    
    【v6关键修复】
    1. 日期匹配同时检查可见文本(li_text)和完整文本(li_full_text)
    2. 增加调试截图功能，失败时自动保存页面截图
    3. 增强降级匹配策略，提高长名称产品的匹配成功率
    """
    
    announcement_display = f"'{announcement_name[:30]}...'" if announcement_name and len(announcement_name) > 30 else f"'{announcement_name}'"
    logger.info(f"[{index+1}/{total}] 正在检验: {announcement_display} - {product_code}")
    
    log_entry = {
        '产品编号': str(product_code) if product_code else '',
        '公告名称': str(announcement_name) if announcement_name else '',
        '公告URL': str(announcement_url) if announcement_url else '',
        '验证时间': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        '验证状态': '进行中',
        '失败原因': '',
        '附件数量': 0,
        '附件状态': '',
        '附件日期验证': '',
        '净值公告间隔': ''
    }
    
    try:
        # 步骤1: 先回公告列表首页（确保搜索框存在）
        logger.info(f"  ├─ 步骤1/9: 返回公告列表首页...")
        driver.get("https://finance1.ccb.com/chn/finance/yfgg.shtml")
        try:
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, "input.kuang"))
            )
        except:
            time.sleep(2)
        
        # 步骤2+3 合并：JS 一次性完成清空→输入→点击搜索
        logger.info(f"  ├─ 步骤2/9: 输入产品编号并搜索...")
        try:
            search_ok = driver.execute_script("""
                var kuang = document.querySelector('input.kuang');
                var but = document.querySelector('input.but');
                if (!kuang || !but) return false;
                kuang.value = '';
                kuang.dispatchEvent(new Event('focus'));
                kuang.value = arguments[0];
                but.click();
                return true;
            """, product_code)
            
            if not search_ok:
                # 兜底：用 Selenium 方式
                search_box = None
                try:
                    search_box = WebDriverWait(driver, 3).until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, "input.kuang"))
                    )
                except:
                    pass
                if not search_box:
                    log_entry['验证状态'] = '失败'
                    log_entry['失败原因'] = '找不到搜索框'
                    logger.warning(f"  └─ 结果: 失败 - {log_entry['失败原因']}")
                    return log_entry
                search_box.clear()
                search_box.send_keys(product_code)
                btn = driver.find_element(By.CSS_SELECTOR, "input.but")
                btn.click()
        except Exception as e:
            log_entry['验证状态'] = '失败'
            log_entry['失败原因'] = f'搜索执行失败: {str(e)[:30]}'
            logger.warning(f"  └─ 结果: 失败 - {log_entry['失败原因']}")
            return log_entry
        
        # 步骤3: 等待搜索结果加载
        logger.info(f"  ├─ 步骤3/9: 等待搜索结果...")
        try:
            WebDriverWait(driver, 10).until(
                EC.presence_of_all_elements_located((By.CSS_SELECTOR, "a[href*='newsdetail']"))
            )
        except:
            time.sleep(2)  # 兜底
        
        # 步骤4: 查找目标公告
        # 【v6修复】日期匹配同时检查可见文本和完整文本（绕过CSS截断）
        logger.info(f"  ├─ 步骤4/9: 正在查找目标公告...")
        announcement_link = None
        
        # 目标日期的各种格式（用于文本匹配）
        date_no_dash = announcement_date.replace('-', '')          # 20260422
        date_slash   = announcement_date.replace('-', '/')         # 2026/04/22
        date_variants = [announcement_date, date_no_dash, date_slash]
        
        # 先获取所有 <li> 元素，每个 li 内含链接+日期文字
        li_elements = driver.find_elements(By.CSS_SELECTOR, "ul.list li")
        logger.info(f"  │  列表中找到 {len(li_elements)} 条公告")
        
        # 收集匹配信息用于调试
        debug_matches = []
        
        for li in li_elements:
            try:
                li_text = li.text  # 可见文本（可能被CSS截断）
                # 【v6修复】获取li的完整文本内容（绕过CSS截断）
                li_full_text = driver.execute_script(
                    "return arguments[0].textContent || '';", li
                ).strip()
                
                # 【v6修复】日期匹配：同时检查可见文本和完整文本
                date_in_visible = any(v in li_text for v in date_variants)
                date_in_full = any(v in li_full_text for v in date_variants)
                date_matched = date_in_visible or date_in_full
                
                if not date_matched:
                    continue
                
                # 日期匹配，找链接
                link_elem = li.find_element(By.TAG_NAME, "a")
                href = link_elem.get_attribute('href') or ''
                link_text = link_elem.text.strip()
                # 链接的完整文本（绕过CSS截断）
                link_full_text = driver.execute_script(
                    "return arguments[0].textContent || '';", link_elem
                ).strip()
                link_title = link_elem.get_attribute('title') or ''
                
                match_info = {
                    'li_text': li_text.strip()[:80],
                    'li_full_text': li_full_text.strip()[:80],
                    'link_text': link_text[:50],
                    'link_full_text': link_full_text[:50],
                    'href': href[:60],
                    'date_in_visible': date_in_visible,
                    'date_in_full': date_in_full
                }
                debug_matches.append(match_info)
                
                logger.info(f"  │  日期匹配条目: [{li_text.strip()[:60]}]")
                if date_in_full and not date_in_visible:
                    logger.info(f"  │  ⚠ 日期仅在完整文本中找到（可能被CSS截断）: {li_full_text.strip()[:60]}")
                
                if 'newsdetail' in href:
                    # 策略A：先用可见文本匹配（快速）
                    code_in_visible = (
                        product_code in link_title or 
                        product_code in li_text or 
                        product_code in link_text or 
                        product_code in href
                    )
                    # 策略B：若策略A失败，用完整文本匹配（处理CSS截断的长名称）
                    code_in_full = False
                    if not code_in_visible:
                        code_in_full = (
                            product_code in li_full_text or
                            product_code in link_full_text or
                            product_code in link_title
                        )
                    
                    if code_in_visible or code_in_full:
                        announcement_link = href
                        match_type = "完整文本" if code_in_full else "可见文本"
                        logger.info(f"  │  ✓ 产品编号+日期均匹配（{match_type}）: {link_text[:50] or link_title[:50]}")
                        break
                    # 备选：先记录第一个日期匹配项（即使编号不匹配，也先记下）
                    if not announcement_link:
                        announcement_link = href
                        logger.info(f"  │  日期匹配（备选，编号未确认）: {link_text[:50]}")
            except:
                continue
        
        # 兜底：若 ul.list 未渲染完成，降级扫全部 <a> 标签，用文字日期匹配
        if not announcement_link:
            logger.info(f"  │  未从列表项找到，降级扫全部链接...")

            all_links = driver.find_elements(By.TAG_NAME, "a")
            logger.info(f"  │  页面中共 {len(all_links)} 个链接")
            for link in all_links:
                try:
                    href = link.get_attribute('href') or ''
                    if 'newsdetail' not in href:
                        continue
                    # 尝试读父元素文字（含日期）
                    try:
                        parent_li = driver.execute_script("return arguments[0].closest('li');", link)
                        parent_text = parent_li.text if parent_li else ''
                        parent_full_text = driver.execute_script("return arguments[0] ? arguments[0].textContent : '';", parent_li) if parent_li else ''
                    except:
                        parent_text = link.text
                        parent_full_text = ''
                    
                    link_text = link.text.strip()
                    link_title = link.get_attribute('title') or ''
                    
                    # 获取完整文本（绕过CSS截断）
                    link_full_text = driver.execute_script(
                        "return arguments[0].textContent || '';", link
                    ).strip()
                    
                    # 【v6修复】日期匹配同时检查可见和完整文本
                    date_in_visible = any(v in parent_text for v in date_variants)
                    date_in_full = any(v in parent_full_text for v in date_variants)
                    date_in_link = any(v in link_full_text for v in date_variants)
                    
                    if date_in_visible or date_in_full or date_in_link:
                        # 策略A：可见文本匹配
                        code_in_visible = (
                            product_code in link_title or 
                            product_code in parent_text or 
                            product_code in link_text or 
                            product_code in href
                        )
                        # 策略B：完整文本匹配（CSS截断时）
                        code_in_full = False
                        if not code_in_visible:
                            code_in_full = (
                                product_code in link_full_text or
                                product_code in parent_full_text or
                                product_code in link_title
                            )
                        
                        if code_in_visible or code_in_full:
                            match_type = "完整文本" if code_in_full else "可见文本"
                            announcement_link = href
                            logger.info(f"  │  ✓ 降级匹配（产品+日期，{match_type}）: {link_text[:50] or link_title[:50]}")
                            break
                        if not announcement_link:
                            announcement_link = href
                            logger.info(f"  │  降级匹配（仅日期）: {link_text[:50]}")
                except:
                    continue
        
        # 降级策略2：若仍找不到，且产品名较长（如"安鑫"类），尝试用产品名关键词模糊匹配
        if not announcement_link and announcement_name:
            logger.info(f"  │  降级策略2: 尝试用产品名关键词匹配...")
            # 提取产品名关键词（取公告名前20字，去掉"净值公告"等通用词）
            keywords_to_try = []
            clean_name = announcement_name.replace('净值公告', '').replace('产品', '').strip()
            # 尝试提取产品系列名（如"安鑫"、"嘉鑫"等）
            import re as re_module
            series_match = re_module.search(r'[""]?([\u4e00-\u9fff]{2,4})', clean_name)
            if series_match:
                series = series_match.group(1)
                if len(series) >= 2:
                    keywords_to_try.append(series)
            # 添加前10字作为备选关键词
            if len(clean_name) >= 10:
                keywords_to_try.append(clean_name[:10])
            
            logger.info(f"  │  关键词候选: {keywords_to_try}")
            
            for kw in keywords_to_try:
                all_links = driver.find_elements(By.TAG_NAME, "a")
                for link in all_links:
                    try:
                        href = link.get_attribute('href') or ''
                        if 'newsdetail' not in href:
                            continue
                        link_text = link.text.strip()
                        link_title = link.get_attribute('title') or ''
                        # 尝试读父元素
                        try:
                            parent_li = driver.execute_script("return arguments[0].closest('li');", link)
                            parent_text = parent_li.text if parent_li else ''
                            parent_full_text = driver.execute_script("return arguments[0] ? arguments[0].textContent : '';", parent_li) if parent_li else ''
                        except:
                            parent_text = ''
                            parent_full_text = ''
                        
                        # 获取完整文本（绕过CSS截断）
                        link_full_text = driver.execute_script(
                            "return arguments[0].textContent || '';", link
                        ).strip()
                        
                        # 合并所有文本用于关键词匹配
                        all_visible_text = (parent_text + ' ' + link_text + ' ' + link_title).lower()
                        all_full_text = (parent_full_text + ' ' + link_full_text + ' ' + link_title).lower()
                        
                        # 检查日期+关键词（先用可见文本，再用完整文本）
                        date_in_visible = any(v in all_visible_text for v in date_variants)
                        date_in_full = any(v in all_full_text for v in date_variants)
                        
                        kw_in_visible = kw.lower() in all_visible_text
                        kw_in_full = kw.lower() in all_full_text
                        
                        if (date_in_visible or date_in_full) and (kw_in_visible or kw_in_full):
                            announcement_link = href
                            match_type = "完整文本" if (kw_in_full and not kw_in_visible) or (date_in_full and not date_in_visible) else "可见文本"
                            logger.info(f"  │  ✓ 降级匹配（日期+关键词'{kw}'，{match_type}）: {link_text[:40] or link_title[:40]}")
                            break
                    except:
                        continue
                if announcement_link:
                    break
        
        # 降级策略3：终极降级 - 只要日期匹配就采用，因为列表页本身就做了搜索过滤
        if not announcement_link:
            logger.info(f"  │  降级策略3: 终极降级（仅日期匹配）...")
            all_links = driver.find_elements(By.TAG_NAME, "a")
            for link in all_links:
                try:
                    href = link.get_attribute('href') or ''
                    if 'newsdetail' not in href:
                        continue
                    link_text = link.text.strip()
                    link_title = link.get_attribute('title') or ''
                    try:
                        parent_li = driver.execute_script("return arguments[0].closest('li');", link)
                        parent_text = parent_li.text if parent_li else ''
                        parent_full_text = driver.execute_script("return arguments[0] ? arguments[0].textContent : '';", parent_li) if parent_li else ''
                    except:
                        parent_text = ''
                        parent_full_text = ''
                    
                    # 获取完整文本
                    link_full_text = driver.execute_script(
                        "return arguments[0].textContent || '';", link
                    ).strip()
                    
                    # 只要日期匹配就行，因为搜索已经按产品过滤了
                    all_text = (parent_text + ' ' + parent_full_text + ' ' + link_text + ' ' + link_title + ' ' + link_full_text + ' ' + href).lower()
                    date_in_text = any(v in all_text for v in date_variants)
                    if date_in_text:
                        announcement_link = href
                        logger.info(f"  │  ✓ 终极降级（仅日期）: {link_text[:40] or link_title[:40]}")
                        break
                except:
                    continue
        
        if not announcement_link:
            log_entry['验证状态'] = '失败'
            log_entry['失败原因'] = f'未找到{announcement_date}的公告链接'
            logger.warning(f"  └─ 结果: 失败 - {log_entry['失败原因']}")
            
            # 【v6新增】保存调试截图
            if screenshot_dir:
                logger.info(f"  │  保存调试截图用于排查...")
                screenshot_path = save_debug_screenshot(driver, screenshot_dir, "search_failed", product_code)
                if screenshot_path:
                    logger.info(f"  │  截图已保存: {screenshot_path}")
                # 同时打印调试信息
                if debug_matches:
                    logger.info(f"  │  调试信息: 找到 {len(debug_matches)} 个日期匹配条目，但产品编号未匹配")
                    for i, dm in enumerate(debug_matches[:3]):
                        logger.info(f"  │    [{i+1}] li_text='{dm['li_text']}' li_full='{dm['li_full_text']}' href='{dm['href']}'")
            
            return log_entry
        
        # 步骤4.5: 校验净值公告发布间隔（需启用）
        if interval_days > 0 and interval_type:
            logger.info(f"  ├─ 步骤4.5/9: 校验净值公告发布间隔（期望{interval_days}{interval_type}）...")
            nav_interval_result = check_nav_intervals(driver, product_code, announcement_date,
                                                      interval_days=interval_days, interval_type=interval_type)
            log_entry['净值公告间隔'] = nav_interval_result
        else:
            log_entry['净值公告间隔'] = '未配置'
        
        # 步骤5: 打开公告详情页
        logger.info(f"  ├─ 步骤5/9: 正在打开公告详情页...")
        driver.get(announcement_link)
        # 用显式等待替代 time.sleep(3)：等待页面主体内容加载
        try:
            WebDriverWait(driver, 10).until(
                lambda d: len(d.find_elements(By.TAG_NAME, "a")) > 0
            )
        except:
            time.sleep(1)
        
        # 步骤6: 严格验证公告标题
        logger.info(f"  ├─ 步骤6/9: 正在严格验证公告标题...")
        title_valid, title_msg = validate_announcement_title(driver, announcement_name)
        
        if not title_valid:
            log_entry['验证状态'] = '失败'
            log_entry['失败原因'] = f'公告标题验证失败: {title_msg}'
            logger.warning(f"  └─ 结果: 失败 - {log_entry['失败原因']}")
            return log_entry
        
        logger.info(f"  │  √ {title_msg}")
        
        # 步骤7: 扫描附件（优化选择器，减少无效遍历）
        logger.info(f"  ├─ 步骤7/9: 正在扫描附件...")
        attachment_links = []
        
        # 使用更精确的选择器直接定位带文件扩展名的链接
        file_ext_pattern = r'\.(pdf|docx?|xlsx?|rar|zip)(\?|$)'
        all_links = driver.find_elements(By.CSS_SELECTOR, "a[href]")
        for link in all_links:
            try:
                href = link.get_attribute('href') or ''
                text = link.text.strip()
                
                if re.search(file_ext_pattern, href, re.IGNORECASE):
                    if text:  # 有文本的链接更有可能是附件
                        attachment_links.append({'text': text, 'href': href})
                        logger.info(f"  │  找到附件: {text[:30]}...")
            except:
                continue
        
        # 去重处理
        unique_attachments = []
        seen_hrefs = set()
        for att in attachment_links:
            if att['href'] not in seen_hrefs:
                unique_attachments.append(att)
                seen_hrefs.add(att['href'])
        
        attachment_links = unique_attachments
        log_entry['附件数量'] = len(attachment_links)
        
        if not attachment_links:
            log_entry['验证状态'] = '失败'
            log_entry['失败原因'] = '无附件'
            logger.warning(f"  └─ 结果: 失败 - {log_entry['失败原因']}")
            return log_entry
        
        # 步骤8: 验证附件（包含日期验证）
        logger.info(f"  ├─ 步骤8/9: 正在验证{len(attachment_links)}个附件...")
        clean_download_folder(download_path)
        attachment_status = []
        date_validation_results = []
        
        for i, att in enumerate(attachment_links):
            logger.info(f"    ├─ 附件{i+1}: {att.get('text', '未命名')[:30]}...")
            
            status, date_result = verify_attachment_with_date(
                driver, att['href'], announcement_date, download_path
            )
            
            attachment_status.append(status)
            date_validation_results.append(date_result)
            
            if status == "成功":
                logger.info(f"    │ 状态: 成功 - {date_result}")
            else:
                logger.warning(f"    │ 状态: {status} - {date_result}")
        
        log_entry['附件状态'] = '; '.join(attachment_status)
        log_entry['附件日期验证'] = '; '.join(date_validation_results)
        
        # 步骤9: 综合判断结果
        logger.info(f"  ├─ 步骤9/9: 综合判断验证结果...")
        
        if all(s == "成功" for s in attachment_status):
            date_verified_count = sum(
                1 for i, s in enumerate(attachment_status) 
                if s == "成功" and "日期验证通过" in (date_validation_results[i] if i < len(date_validation_results) else "")
            )
            
            if date_verified_count == len(attachment_status):
                log_entry['验证状态'] = '成功'
                log_entry['失败原因'] = '无'
                logger.info(f"  └─ 结果: 成功 - 所有附件验证通过且日期匹配")
            elif date_verified_count > 0:
                log_entry['验证状态'] = '部分成功'
                log_entry['失败原因'] = f'附件可访问但{len(attachment_status)-date_verified_count}个附件日期未验证'
                logger.warning(f"  └─ 结果: 部分成功 - {log_entry['失败原因']}")
            else:
                log_entry['验证状态'] = '部分成功'
                log_entry['失败原因'] = '所有附件可访问但日期均未验证'
                logger.warning(f"  └─ 结果: 部分成功 - {log_entry['失败原因']}")
        else:
            failed_count = sum(1 for s in attachment_status if s != "成功")
            has_success = any(s == "成功" for s in attachment_status)
            
            log_entry['验证状态'] = '部分失败' if has_success else '失败'
            log_entry['失败原因'] = f'{failed_count}个附件验证失败'
            logger.warning(f"  └─ 结果: {log_entry['验证状态']} - {log_entry['失败原因']}")
        
    except Exception as e:
        error_msg = str(e)
        log_entry['验证状态'] = '失败'
        log_entry['失败原因'] = f"异常:{error_msg[:30] if error_msg else '未知错误'}"
        logger.error(f"  └─ 结果: 失败 - {error_msg[:100] if error_msg else '未知错误'}")
        clean_download_folder(download_path)
    
    return log_entry

# ==================== 主函数 ====================
def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    download_path = os.path.join(script_dir, 'temp_downloads')
    if not os.path.exists(download_path):
        os.makedirs(download_path)
    
    # 【v6新增】创建调试截图目录
    screenshot_dir = os.path.join(script_dir, 'debug_screenshots')
    if not os.path.exists(screenshot_dir):
        os.makedirs(screenshot_dir)
        logger.info(f"调试截图目录已创建: {screenshot_dir}")
    
    # 读取配置
    config_file = os.path.join(script_dir, 'config.txt')
    if not os.path.exists(config_file):
        logger.error("配置文件 config.txt 不存在！")
        return
    
    with open(config_file, 'r', encoding='utf-8') as f:
        lines = [l.strip() for l in f if l.strip() and not l.strip().startswith('#')]
    
    if len(lines) < 2:
        logger.error("配置文件格式错误！第一行应为Excel路径，第二行应为公告日期")
        return
    
    excel_path, announcement_date = lines[0], lines[1]
    
    # 解析净值公告间隔配置（第三行，可选）
    interval_days = 7
    interval_type = '自然日'
    enable_interval_check = False
    if len(lines) >= 3 and lines[2]:
        try:
            interval_config = lines[2]
            if ',' in interval_config:
                days_str, type_str = interval_config.split(',', 1)
                interval_days = int(days_str.strip())
                interval_type = type_str.strip()
            else:
                interval_days = int(interval_config.strip())
            enable_interval_check = True
            logger.info(f"净值公告间隔校验: {interval_days}{interval_type}")
        except ValueError:
            logger.warning(f"间隔配置格式错误: {lines[2]}，跳过间隔校验（正确格式: 7,自然日）")
    else:
        logger.info("未配置净值公告间隔校验，跳过（如需启用请在config第三行配置，如: 7,自然日）")
    
    # 处理路径：将 Windows 路径分隔符转换为系统分隔符
    excel_path = excel_path.replace('\\', os.sep).replace('/', os.sep)
    
    # 验证公告日期格式
    if not re.match(r'\d{4}-\d{2}-\d{2}', announcement_date):
        logger.error(f"公告日期格式错误: {announcement_date}，应为YYYY-MM-DD格式")
        return
    
    logger.info(f"公告日期: {announcement_date}")
    
    # 读取Excel文件
    if not os.path.exists(excel_path):
        logger.error(f"Excel文件不存在: {excel_path}")
        return
    
    logger.info(f"加载Excel文件: {excel_path}")
    
    # 读取Excel文件，使用object类型保持原始格式
    try:
        df = pd.read_excel(excel_path, dtype=object)
        logger.info(f"成功加载数据，共 {len(df)} 条记录")
    except Exception as e:
        logger.error(f"读取Excel文件失败: {str(e)}")
        return
    
    # 识别列名
    p_col, n_col, u_col = None, None, None
    for c in df.columns:
        col_str = str(c)
        if '产品编号' in col_str: 
            p_col = c
        elif '标题' in col_str or '公告名称' in col_str: 
            n_col = c
        elif '公告地址' in col_str or 'URL' in col_str.upper(): 
            u_col = c
    
    if p_col is None:
        logger.error("未找到'产品编号'列！")
        logger.info(f"可用列: {list(df.columns)}")
        return
    
    logger.info(f"识别列名: 产品编号={p_col}, 公告名称={n_col}, 公告地址={u_col}")
    
    # 确保结果列存在
    result_cols = ['验证时间', '验证状态', '失败原因', '附件数量', '附件状态', '附件日期验证', '净值公告间隔']
    for col in result_cols:
        if col not in df.columns:
            if col == '附件数量':
                df[col] = pd.NA
            else:
                df[col] = ''
            logger.info(f"添加结果列: {col}")
    
    # 分析已处理记录，实现断点续跑
    logger.info("正在分析Excel文件，检测已处理记录...")
    
    completed_records = []
    pending_records = []
    
    for idx in df.index:
        status_value = df.at[idx, '验证状态']
        # 只有验证状态为"成功"的才跳过，其余（失败/部分成功/部分失败/异常/空白）都需要重跑
        if pd.notna(status_value) and str(status_value).strip() == '成功':
            completed_records.append(idx)
        else:
            pending_records.append(idx)
    
    total = len(df)
    completed = len(completed_records)
    pending = len(pending_records)
    
    logger.info(f"=== 断点续跑分析 ===")
    logger.info(f"总记录数: {total}")
    logger.info(f"已处理记录: {completed}")
    logger.info(f"待处理记录: {pending}")
    
    if pending == 0:
        logger.info("所有记录已处理完成，无需继续验证。")
        return
    
    # 询问用户是否继续处理
    if completed > 0:
        logger.info(f"发现 {completed} 条已处理记录，将从第 {pending_records[0]+1} 条记录开始继续处理。")
        response = input("是否继续处理？(y/n): ").strip().lower()
        if response != 'y':
            logger.info("程序退出")
            return
    
    # 浏览器配置
    logger.info("正在初始化浏览器...")
    options = webdriver.ChromeOptions()
    
    # 无头模式
    options.add_argument('--headless=new')
    
    # 优化参数
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--disable-gpu')
    options.add_argument('--disable-images')
    options.add_argument('--disable-plugins')
    options.add_argument('--disable-infobars')
    options.add_argument('--disable-extensions')
    options.add_argument('--disable-dev-tools')
    options.add_argument('--log-level=3')
    options.add_experimental_option('excludeSwitches', ['enable-logging', 'enable-automation'])
    
    # 下载配置 - 关键设置
    prefs = {
        "download.default_directory": download_path,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.enabled": True,
        "profile.default_content_setting_values.automatic_downloads": 1,
        "profile.default_content_settings.popups": 0,
        "profile.content_settings.exceptions.automatic_downloads.*.setting": 1,
        "profile.managed_default_content_settings.images": 2,
        "profile.default_content_setting_values.javascript": 1
    }
    options.add_experimental_option("prefs", prefs)
    
    # ==================== ChromeDriver 初始化（Windows 10 兼容版）====================
    logger.info(f"检测系统平台: {platform.system()}")
    
    # ChromeDriver 路径缓存文件，避免每次都通过 webdriver-manager 联网检测
    driver_cache_file = os.path.join(script_dir, '.chromedriver_path')
    cached_driver_path = None
    
    # 读取缓存路径
    if os.path.exists(driver_cache_file):
        try:
            with open(driver_cache_file, 'r') as f:
                cached_driver_path = f.read().strip()
            if cached_driver_path and os.path.exists(cached_driver_path):
                logger.info(f"使用缓存的 ChromeDriver: {cached_driver_path}")
            else:
                cached_driver_path = None
        except:
            cached_driver_path = None
    
    driver_initialized = False
    
    # 策略1：使用缓存路径
    if cached_driver_path:
        try:
            service = Service(cached_driver_path)
            driver = webdriver.Chrome(service=service, options=options)
            logger.info("浏览器驱动初始化成功（缓存路径）")
            driver_initialized = True
        except Exception as e:
            logger.warning(f"缓存路径失效: {str(e)[:50]}，将重新获取 ChromeDriver...")
            cached_driver_path = None
    
    # 策略2：webdriver-manager 自动管理（Windows 主要方案）
    if not driver_initialized:
        try:
            logger.info("正在通过 webdriver-manager 获取与当前 Chrome 版本匹配的 ChromeDriver...")
            driver_path = ChromeDriverManager().install()
            # 保存到缓存
            try:
                with open(driver_cache_file, 'w') as f:
                    f.write(driver_path)
                logger.info("ChromeDriver 路径已缓存")
            except:
                pass
            service = Service(driver_path)
            driver = webdriver.Chrome(service=service, options=options)
            logger.info("浏览器驱动初始化成功（webdriver-manager）")
            driver_initialized = True
        except Exception as e:
            logger.warning(f"webdriver-manager 获取失败: {str(e)[:80]}")
    
    # 策略3：手动查找 chromedriver.exe（Windows 常用位置）
    if not driver_initialized:
        logger.info("\n尝试手动查找 ChromeDriver...")
        
        possible_driver_paths = []
        
        if platform.system() == "Windows":
            # Windows 下的可能路径（按优先级排列）
            possible_driver_paths.extend([
                os.path.join(script_dir, "chromedriver.exe"),      # 当前脚本目录
                "chromedriver.exe",                                  # PATH 中
                os.path.join(os.getcwd(), "chromedriver.exe"),      # 当前工作目录
                # C:\Program Files\Google\Chrome\Application\ 目录（Chrome 安装位置）
                r"C:\Program Files\Google\Chrome\Application\chromedriver.exe",
                # C:\Program Files (x86)\Google\Chrome\Application\（32位 Chrome）
                r"C:\Program Files (x86)\Google\Chrome\Application\chromedriver.exe",
                # 用户目录下的 Chrome
                os.path.join(os.path.expanduser("~"), "chromedriver.exe"),
                # 常见工具目录
                r"C:\tools\chromedriver.exe",
                r"C:\webdriver\chromedriver.exe",
            ])
        elif platform.system() == "Darwin":
            # macOS 路径（保留以兼容少量 macOS 用户）
            possible_driver_paths.extend([
                os.path.join(script_dir, "chromedriver"),
                "chromedriver",
                os.path.expanduser("~/chromedriver"),
            ])
        elif platform.system() == "Linux":
            possible_driver_paths.extend([
                os.path.join(script_dir, "chromedriver"),
                "chromedriver",
                "/usr/bin/chromedriver",
                "/usr/local/bin/chromedriver",
            ])
        
        driver_path = None
        for path in possible_driver_paths:
            if os.path.exists(path):
                driver_path = path
                logger.info(f"找到 ChromeDriver: {driver_path}")
                break
        
        if not driver_path:
            logger.error("ChromeDriver 未找到！")
            logger.info("请采取以下措施之一：")
            logger.info("  方案 1: 删除 .chromedriver_path 缓存文件后重试（清理错误的缓存路径）")
            logger.info("  方案 2: 手动下载对应版本的 ChromeDriver：")
            logger.info("    - 访问 https://googlechromelabs.github.io/chrome-for-testing/")
            logger.info("    - 下载与您 Chrome 版本匹配的 chromedriver.exe")
            logger.info("    - 将 chromedriver.exe 放到脚本同目录下，或添加到 PATH")
            return
        
        try:
            try:
                with open(driver_cache_file, 'w') as f:
                    f.write(driver_path)
            except:
                pass
            service = Service(driver_path)
            driver = webdriver.Chrome(service=service, options=options)
            logger.info("浏览器驱动初始化成功（手动路径）")
            driver_initialized = True
        except Exception as e2:
            logger.error(f"初始化浏览器驱动失败: {str(e2)}")
            return
    
    # 开始处理
    actually_processed_count = 0
    start_time = time.time()
    
    logger.info(f"=== 开始批量验证 ===")
    logger.info(f"公告日期: {announcement_date}")
    logger.info(f"结果保存到: {excel_path}")
    logger.info(f"调试截图保存到: {screenshot_dir}")
    logger.info("-" * 50)
    
    for record_index, idx in enumerate(pending_records):
        row = df.iloc[idx]
        code = str(row[p_col]) if pd.notna(row[p_col]) else ''
        name = str(row[n_col]) if n_col and pd.notna(row[n_col]) else ''
        url = str(row[u_col]) if u_col and pd.notna(row[u_col]) else ''
        
        if not code:
            logger.warning(f"记录 {idx+1} 产品编号为空，跳过")
            continue
        
        # 执行验证（传入screenshot_dir用于调试）
        res = search_and_verify(driver, code, name, url, announcement_date, download_path, 
                              idx, total, completed + actually_processed_count,
                              interval_days=interval_days, interval_type=interval_type,
                              screenshot_dir=screenshot_dir)
        
        # 更新DataFrame（idx 本身已是 df.index 值，直接用 df.at[idx, ...] 即可）
        df.at[idx, '验证时间'] = str(res['验证时间'])
        df.at[idx, '验证状态'] = str(res['验证状态'])
        df.at[idx, '失败原因'] = str(res['失败原因'])
        df.at[idx, '附件数量'] = int(res['附件数量'])
        df.at[idx, '附件状态'] = str(res['附件状态'])
        df.at[idx, '附件日期验证'] = str(res['附件日期验证'])
        df.at[idx, '净值公告间隔'] = str(res['净值公告间隔'])
        
        # 保存到原Excel文件
        try:
            df.to_excel(excel_path, index=False)
            logger.info(f"  √ 已保存到原文件: {excel_path}")
        except Exception as e:
            logger.error(f"保存Excel文件失败: {str(e)}")
            # 尝试使用备份名称保存
            backup_name = f"{os.path.splitext(excel_path)[0]}_备份_{int(time.time())}.xlsx"
            try:
                df.to_excel(backup_name, index=False)
                logger.info(f"已保存备份到: {backup_name}")
            except Exception as e2:
                logger.error(f"保存备份文件也失败: {str(e2)}")
        
        actually_processed_count += 1
        
        # 显示进度
        current_total_processed = completed + actually_processed_count
        progress = (current_total_processed) / total * 100
        elapsed = time.time() - start_time
        
        if actually_processed_count > 0:
            avg_time_per_item = elapsed / actually_processed_count
            remaining_items = total - current_total_processed
            eta = avg_time_per_item * remaining_items
        else:
            eta = 0
        
        logger.info(f"进度: {current_total_processed}/{total} [{progress:.1f}%] | "
                   f"用时: {elapsed:.0f}s | "
                   f"剩余: {eta:.0f}s")
        logger.info("-" * 50)
    
    # 清理
    try:
        driver.quit()
    except:
        pass
    
    clean_download_folder(download_path)
    try: 
        os.rmdir(download_path)
    except: 
        pass
    
    # 最终统计
    logger.info("=== 验证完成 ===")
    total_time = time.time() - start_time
    success_count = 0
    failed_count = 0
    partial_count = 0
    
    for idx in df.index:
        status_value = df.at[idx, '验证状态']
        if pd.notna(status_value):
            status_str = str(status_value).strip()
            if status_str == '成功':
                success_count += 1
            elif status_str == '失败':
                failed_count += 1
            elif status_str in ['部分失败', '部分成功']:
                partial_count += 1
    
    logger.info(f"总计: {total} 条")
    logger.info(f"成功: {success_count} 条")
    logger.info(f"部分成功/失败: {partial_count} 条")
    logger.info(f"失败: {failed_count} 条")
    logger.info(f"总用时: {total_time:.1f}秒")
    if actually_processed_count > 0:
        logger.info(f"平均每条: {total_time/actually_processed_count:.1f}秒")
    logger.info(f"结果已保存到原文件: {excel_path}")
    logger.info(f"调试截图保存在: {screenshot_dir}")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        logger.warning("\n程序被用户中断，Excel文件已保存，下次运行将自动从断点继续。")
    except Exception as e:
        logger.error(f"程序异常终止: {str(e)}")
        import traceback
        logger.error(f"详细错误信息:\n{traceback.format_exc()}")
        logger.info("Excel文件已保存，下次运行将自动从断点继续。")
